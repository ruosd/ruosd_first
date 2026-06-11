import os
import logging
import re
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from pathlib import Path
from enum import Enum
import pypdf
from docx import Document
import tiktoken

logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """文档类型枚举"""
    PDF = "pdf"
    TXT = "txt"
    DOCX = "docx"
    UNKNOWN = "unknown"


@dataclass
class DocumentChunk:
    """文档块 - 小块"""
    chunk_id: str
    content: str
    chunk_index: int
    parent_id: str
    metadata: Dict[str, Any]


@dataclass
class DocumentSection:
    """文档段 - 大块"""
    section_id: str
    content: str
    section_title: str
    section_index: int
    chunks: List[DocumentChunk]
    metadata: Dict[str, Any]


@dataclass
class ProcessedDocument:
    """处理后的完整文档"""
    doc_id: str
    file_name: str
    file_type: DocumentType
    sections: List[DocumentSection]
    metadata: Dict[str, Any]


class DocumentLoader:
    """文档加载器"""
    
    @staticmethod
    def detect_type(file_path: str) -> DocumentType:
        """检测文档类型"""
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            return DocumentType.PDF
        elif ext == ".txt":
            return DocumentType.TXT
        elif ext == ".docx":
            return DocumentType.DOCX
        return DocumentType.UNKNOWN
    
    @staticmethod
    def load_pdf(file_path: str) -> str:
        """加载PDF文档 — 优先 pdfplumber（中文支持好），失败回退 pypdf"""
        text = ""

        # 方案 A: pdfplumber（中文兼容性好）
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n\n"
            if text.strip():
                logger.info(f"PDF(pdfplumber)加载成功: {file_path}, {len(pdf.pages)} 页")
                return text
        except Exception as e:
            logger.warning(f"pdfplumber 加载失败: {e}，回退 pypdf")

        # 方案 B: pypdf 兜底
        try:
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n\n"
            logger.info(f"PDF(pypdf)加载成功: {file_path}, {len(reader.pages)} 页")
        except Exception as e:
            logger.error(f"PDF加载失败: {e}")

        return text
    
    @staticmethod
    def load_txt(file_path: str) -> str:
        """加载TXT文档"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            logger.info(f"TXT加载成功: {file_path}")
        except Exception as e:
            logger.error(f"TXT加载失败: {e}")
            text = ""
        return text
    
    @staticmethod
    def load_docx(file_path: str) -> str:
        """加载DOCX文档"""
        try:
            doc = Document(file_path)
            text = "\n\n".join([para.text for para in doc.paragraphs])
            logger.info(f"DOCX加载成功: {file_path}")
        except Exception as e:
            logger.error(f"DOCX加载失败: {e}")
            text = ""
        return text
    
    @classmethod
    def load(cls, file_path: str) -> Tuple[str, DocumentType]:
        """加载文档"""
        doc_type = cls.detect_type(file_path)
        
        if doc_type == DocumentType.PDF:
            text = cls.load_pdf(file_path)
        elif doc_type == DocumentType.TXT:
            text = cls.load_txt(file_path)
        elif doc_type == DocumentType.DOCX:
            text = cls.load_docx(file_path)
        else:
            logger.warning(f"不支持的文档类型: {file_path}")
            text = ""
        
        return text, doc_type


class TextSplitter:
    """文本分割器 - 支持分层切割策略"""
    
    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 50,
        section_size: int = 2048
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.section_size = section_size
        try:
            self.tokenizer = tiktoken.get_encoding("cl100k_base")
        except Exception:
            self.tokenizer = None
    
    def _token_count(self, text: str) -> int:
        """计算token数"""
        if self.tokenizer:
            return len(self.tokenizer.encode(text))
        return len(text) // 2  # 简单估算
    
    def _split_by_paragraphs(self, text: str) -> List[str]:
        """按自然段落分割"""
        paragraphs = re.split(r'\n\s*\n', text.strip())
        return [p.strip() for p in paragraphs if p.strip()]
    
    def _detect_headings(self, text: str) -> List[Tuple[str, int, int]]:
        """检测标题（适合Word文档）"""
        headings = []
        # 常见标题模式
        heading_patterns = [
            r'^第[一二三四五六七八九十百千万\d]+[章节条]',
            r'^[一二三四五六七八九十]+[、.]',
            r'^\d+[\.\d]*[、.]',
            r'^[A-Z][A-Z\s]+[.:]?$',
            r'^#+\s+',
        ]
        
        paragraphs = self._split_by_paragraphs(text)
        start_idx = 0
        
        for i, para in enumerate(paragraphs):
            for pattern in heading_patterns:
                if re.match(pattern, para.strip()):
                    headings.append((para.strip(), start_idx, i))
                    break
            start_idx += len(para) + 2
        
        return headings
    
    def _split_into_sections(
        self,
        text: str,
        doc_type: DocumentType
    ) -> List[Tuple[str, str]]:
        """分割成大块（sections）"""
        sections = []
        
        if doc_type == DocumentType.DOCX:
            # Word文档按标题分割
            paragraphs = self._split_by_paragraphs(text)
            headings = self._detect_headings(text)
            
            if headings:
                current_title = "文档开头"
                current_content = []
                
                for i, para in enumerate(paragraphs):
                    is_heading = any(h[2] == i for h in headings)
                    if is_heading:
                        if current_content:
                            sections.append((current_title, "\n\n".join(current_content)))
                        current_title = para.strip()
                        current_content = []
                    else:
                        current_content.append(para)
                
                if current_content:
                    sections.append((current_title, "\n\n".join(current_content)))
            else:
                sections.append(("文档内容", text))
        else:
            # 其他文档按自然段落和大小分割
            sections = self._split_by_content_size(text)
        
        return sections
    
    def _split_by_content_size(self, text: str) -> List[Tuple[str, str]]:
        """按内容大小分割成大块"""
        paragraphs = self._split_by_paragraphs(text)
        sections = []
        current_section = []
        current_tokens = 0
        
        for para in paragraphs:
            para_tokens = self._token_count(para)
            
            if current_tokens + para_tokens > self.section_size and current_section:
                title = f"章节 {len(sections) + 1}"
                sections.append((title, "\n\n".join(current_section)))
                current_section = [para]
                current_tokens = para_tokens
            else:
                current_section.append(para)
                current_tokens += para_tokens
        
        if current_section:
            title = f"章节 {len(sections) + 1}"
            sections.append((title, "\n\n".join(current_section)))
        
        return sections
    
    def _split_section_into_chunks(
        self,
        section_content: str,
        section_title: str
    ) -> List[DocumentChunk]:
        """将大块分割成小块"""
        chunks = []
        sentences = re.split(r'[。！？.!?\n]+', section_content)
        current_chunk = []
        current_tokens = 0
        
        for i, sentence in enumerate(sentences):
            sentence = sentence.strip()
            if not sentence:
                continue
            
            sentence_tokens = self._token_count(sentence)
            
            if current_tokens + sentence_tokens > self.chunk_size and current_chunk:
                chunk_content = "".join(current_chunk)
                chunk = DocumentChunk(
                    chunk_id=f"chunk_{len(chunks)}",
                    content=chunk_content,
                    chunk_index=len(chunks),
                    parent_id="",
                    metadata={
                        "section_title": section_title
                    }
                )
                chunks.append(chunk)
                current_chunk = [sentence]
                current_tokens = sentence_tokens
            else:
                current_chunk.append(sentence)
                current_tokens += sentence_tokens
        
        if current_chunk:
            chunk_content = "".join(current_chunk)
            chunk = DocumentChunk(
                chunk_id=f"chunk_{len(chunks)}",
                content=chunk_content,
                chunk_index=len(chunks),
                parent_id="",
                metadata={
                    "section_title": section_title
                }
            )
            chunks.append(chunk)
        
        return chunks
    
    def split(
        self,
        text: str,
        doc_type: DocumentType,
        doc_id: str
    ) -> ProcessedDocument:
        """
        分层切割文档
        
        Args:
            text: 文档文本
            doc_type: 文档类型
            doc_id: 文档ID
            
        Returns:
            处理后的文档对象
        """
        if not text:
            return ProcessedDocument(
                doc_id=doc_id,
                file_name="",
                file_type=doc_type,
                sections=[],
                metadata={}
            )
        
        # 第一步：分割成大块（sections）
        section_data = self._split_into_sections(text, doc_type)
        sections = []
        
        for i, (section_title, section_content) in enumerate(section_data):
            section_id = f"{doc_id}_section_{i}"
            
            # 第二步：将大块分割成小块（chunks）
            chunks = self._split_section_into_chunks(section_content, section_title)
            
            # 设置小块的父ID
            for chunk in chunks:
                chunk.parent_id = section_id
            
            section = DocumentSection(
                section_id=section_id,
                content=section_content,
                section_title=section_title,
                section_index=i,
                chunks=chunks,
                metadata={
                    "chunk_count": len(chunks),
                    "token_count": self._token_count(section_content)
                }
            )
            sections.append(section)
        
        processed = ProcessedDocument(
            doc_id=doc_id,
            file_name="",
            file_type=doc_type,
            sections=sections,
            metadata={
                "total_sections": len(sections),
                "total_chunks": sum(len(s.chunks) for s in sections)
            }
        )
        
        logger.info(f"文档切割完成: {doc_id}, {len(sections)}个section, {processed.metadata['total_chunks']}个chunk")
        return processed


class DocumentProcessor:
    """文档处理器 - 整合加载和切割"""
    
    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 50,
        section_size: int = 2048
    ):
        self.loader = DocumentLoader()
        self.splitter = TextSplitter(chunk_size, chunk_overlap, section_size)
    
    def process_file(
        self,
        file_path: str,
        doc_id: Optional[str] = None
    ) -> ProcessedDocument:
        """处理单个文件"""
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            return ProcessedDocument("", "", DocumentType.UNKNOWN, [], {})
        
        if doc_id is None:
            doc_id = Path(file_path).stem
        
        text, doc_type = self.loader.load(file_path)
        processed = self.splitter.split(text, doc_type, doc_id)
        processed.file_name = Path(file_path).name
        
        return processed
    
    def process_text(
        self,
        text: str,
        doc_id: str,
        doc_type: DocumentType = DocumentType.TXT
    ) -> ProcessedDocument:
        """直接处理文本"""
        processed = self.splitter.split(text, doc_type, doc_id)
        return processed
